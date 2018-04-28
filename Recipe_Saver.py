#! python3
# Scrapes SeriousEats recipes and saves them to a word file
# Note: You must put your own folder file path in the 'document.save' function to choose a save location
# Uses beautifulsoup (bs4) and docx libraries

import urllib
import io
import bs4, requests
from docx import Document
from docx.shared import Inches


# checks if user wants to save a new recipe or end the program
def recipe_saver():
    status = input("Would you like to start saving recipes? Enter 'yes' to save a recipe or 'no' to exit: ")
    # asks if the user would like to start entering url's
    if status == 'no':
        print('Goodbye')
    while status != 'no':  # continues running until user tells program to stop
        status = str(input("Please enter a recipe URL or 'no' to exit:"))  # saves a recipe url from input
        if status != 'no':
            recipe_scraper(status)  # calls recipe scraper function using url
    print('Goodbye')  # exits program


def recipe_scraper(recipe_url):
    recipe_page = requests.get(recipe_url)  # downloads the page contents
    beautifulsoup_recipe = bs4.BeautifulSoup(recipe_page.content, 'html.parser')  # parses to beautifulsoup item
    recipe_name = beautifulsoup_recipe.select('''
    body > div.content-main > section.entry-container > article > header > div > h1''')  # extracts title
    recipe_name = recipe_name[0].text.strip()  #removes extra characters
    recipe_image = beautifulsoup_recipe.findAll('img')  # finds images in recipe
    recipe_image = str(recipe_image[0]['src'])  # gets primary recipe image url
    image_req = urllib.request.Request(recipe_image, headers={'User-Agent': '''Mozilla/5.0 (Windows NT 6.1) 
        AppleWebKit/537.36 (KHTML, like Gecko) Chrome/41.0.2228.0 Safari/537.36'''})  # accesses image
    image = urllib.request.urlopen(image_req)
    image_url = io.BytesIO(image.read())  # opens image locally
    ingredients =  beautifulsoup_recipe.select('''
    body > div:nth-of-type(2) > section.entry-container > article > div > div.recipe-wrapper > div.recipe-ingredients >
     ul''')  # extracts ingredients
    ingredients = ingredients[0].text.strip()
    directions = beautifulsoup_recipe.select('''
    body > div.content-main > section.entry-container > article > div > div.recipe-wrapper > div.recipe-procedures >
     ol''')  # extracts directions
    directions = directions[0].text.strip()
    directions = directions.split('\n\n')  # removes newlines inside string
    for item in directions:
        if item == '':
            directions[directions.index(item)] = '\n'  # inserts newlines between direction items
    directions = ''.join(directions)
    notes = beautifulsoup_recipe.select('#entry-text > div.recipe-introduction-body > p:nth-of-type(2)') #extracts notes
    notes = notes[0].text.strip()
    try:
        extra_notes = beautifulsoup_recipe.select('''body > div:nth-of-type(2) > section.entry-container > article > div > 
    div.recipe-bottom > aside.callout.callout-bottom.callout-bottom-recipe.recipe-notes > span > p''')
        extra_notes = extra_notes[0].text.strip()
    except IndexError:
        extra_notes = ''
    active_time = beautifulsoup_recipe.select('''
    body > div:nth-of-type(2) > section.entry-container > article > div > div.recipe-wrapper > ul > li:nth-of-type(2) 
    > span.info''')  # extracts active time
    active_time = active_time[0].text.strip()
    total_time = beautifulsoup_recipe.select('''
    body > div:nth-of-type(2) > section.entry-container > article > div > div.recipe-wrapper > ul > li:nth-of-type(3) 
    > span.info''')  # extracts total time
    total_time = total_time[0].text.strip()
    recipe_items = [recipe_url, recipe_name, active_time, total_time, image_url, notes, extra_notes, ingredients, directions] # saves items to list
    create_recipe_doc(recipe_items)


# creates, populates, and saves a word document with the scraped information
def create_recipe_doc(recipe_items):
    document = Document()  # opens new word document file
    document.add_heading(recipe_items[1])  # starts populating with scraped information
    document.add_heading('Active Time: ' + recipe_items[2] + '\t' + 'Total Time: ' + recipe_items[3] + '\n', level=2)
    document.add_picture(recipe_items[4], width=Inches(6.0))
    document.add_paragraph()
    document.add_heading('Notes', level=2)
    document.add_paragraph(recipe_items[5] + '\n' + recipe_items[6] + '\n')
    document.add_heading('Ingredients', level=2)
    document.add_paragraph(recipe_items[7] + '\n')
    document.add_heading('Directions', level=2)
    document.add_paragraph(recipe_items[8] + '\n')
    document.add_paragraph('Source: ' + recipe_items[0])  # finishes populating with scraped information
    document.save(r'C:\Users\Ben\Documents\Recipes\Web Recipes\\' + str(recipe_items[1]) + '.docx')
    # saves to a file with recipe name as filename

recipe_saver()
